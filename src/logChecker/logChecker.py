#!/usr/bin/env python3
#
# Copyright (C) 2024 Lucas Aimaretto / laimaretto@gmail.com, Beatriz Bonafe / bonafencb@gmail.com , Kathleen Mendonca / kathleencristine20@gmail.com
# Copyright (C) 2023 Lucas Aimaretto / laimaretto@gmail.com
# Copyright (C) 2020 Manuel Saldivar / manuelsaldivar@outlook.com.ar, Lucas Aimaretto / laimaretto@gmail.com
# 
# This is logChecker
# 
# logChecker is free software: you can redistribute it and/or modify
# it under the terms of the 3-clause BSD License.
# 
# logChecker is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY of any kind whatsoever.
# 

import textfsm
import pandas as pd
import glob
import argparse
from sys import platform as _platform
import json
import re
from ttp import ttp
import os
import io

import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt


DATA_VALUE         = 'Value'
DATA_COMMAND       = '#Command:'
DATA_MAJOR_DWN     = '#majorDown:'
DATA_FLTR_COLS     = '#filterColumns:'
DATA_FLTR_ACTN     = '#filterAction:'
DATA_SHOW_DIFF_COL = '#showDiffColumns'

INDEX_COL = {
	'sheet' : {
		'position': 0, 'col': 'A:A', 'colName': 'Sheet', 'width': 22,
	},
	'command' : {
		'position': 1, 'col': 'B:B', 'colName': 'Command', 'width': 30,
	},
    'status' : {
		'position': 2, 'col': 'C:C', 'colName' : 'Status', 'width' : 20,
	}
}

RTR_ID = dict(
	name = ['NAME'],
	both = ['NAME','IP'],
	ip   = ['IP']
)

CELL_COLOR = 'black'
CELL_FONT_SIZE = '12'
NO_MATCH = ['No matching entries found','No Matching Entries Found']

D_STATUS = dict(
	no_parsing = dict( #Hay comando, no parsing
		colorTab = '#4092FF', #blue
		warnText = '####### CHANGES DETECTED #######',
		errText  = '####### NO Parsing Detected #######',
		shortText = "Can't Parsing",
		),
	no_matching_entries = dict(
		colorTab = '#CCCCCC', #gray
		warnText = '####### CHANGES DETECTED #######',
		errText  = '####### No Matching Entries #######',
		shortText = 'No Matching Entries',
		),
	no_template = dict( #no hay template, no hay parsing
		colorTab = '#9765FE', #purple
		warnText = '####### NO Template in template folder #######',
		errText  = '####### NO Template in template folder #######',
		shortText = 'No Template',
		),
	no_data = dict( #hay comando, hay parsing, pero sin datos
		colorTab = '#F06AE5', #pink
		warnText = '####### No Data in Command #######',
		errText  = '####### No Data in Command #######',
		shortText = 'No Data in Log',
		),
	ok = dict(
		colorTab = '#37CC73', #green
		warnText = '####### NO POST-TASK CHANGES DETECTED #######',
		errText  = '####### NO MAJOR ERRORS FOUND #######',
		shortText = 'Ok!',
		),
	changes_detected = dict(
		colorTab = '#FFD568', #yellow
		warnText = '####### CHANGES DETECTED #######',
		errText  = '####### NO MAJOR ERRORS FOUND #######',
		shortText = 'Warning',
		),
	major_errors = dict(
		colorTab = '#F47F31', #orange
		warnText = '####### CHANGES DETECTED #######',
		errText  = '####### MAJOR ERRORS DETECTED POST-TASK #######',
		shortText = 'Major Errors',
	)
)

GENERAL_TEMPL = 'general.template'

GENERAL_TEMPL_LINES = """#Command: .+
#Timos: any
#Version: 1.0.0
Value Lines (.+)

Start
  ^${Lines} -> Record"""

NON_COMMAND_KEYS = ['name','ip','version','hwType','#FINSCRIPT','exit all','']

def readTemplate(fileTemplate, templateFolder, templateEngine):
	'''
	Read the list of templates passed by CSV of textFSM and return template read list (read)

	List of parsed variable names, list of template names

	If fileTemplate is omitted, then all the templates inside the folder are considered.
	'''

	if fileTemplate != '':
		with open(fileTemplate,'r') as f:
			templates = [x.replace('\n','') for x in f.readlines()]
	else:
		if os.path.exists(templateFolder):
			templates = [f.replace(templateFolder,'') for f in glob.glob(templateFolder + '*') if 'majorFile.yml' not in f]
		else: #Si no hay carpeta con templates...
			templateFolder = ''
			templates = []

	d = {}

	d[GENERAL_TEMPL] = {
	'templateColumns':[],
	'commandKey':'',
	'majorDown':['down','dwn'], #En función findMajor, case=False. Aquí no es necesario tener 'Down' y 'Dwn'
	'filterColumns':[],
	'filterAction':None,
	'showDiffColumns':[]
	}

	templates.append(GENERAL_TEMPL)

	for i,tmpltName in enumerate(templates):

		d[tmpltName] = {
			'templateColumns':[],
			'commandKey':'',
			'majorDown':['down','dwn'], #En función findMajor, case=False. Aquí no es necesario tener 'Down' y 'Dwn'
			'filterColumns':[],
			'filterAction':None,
			'showDiffColumns':[],
		}

		if tmpltName == GENERAL_TEMPL:
			tmpltLines = GENERAL_TEMPL_LINES.splitlines()

		else:
			fName = templateFolder+tmpltName
			try:
				with open(fName) as f:
					tmpltLines = f.readlines()
			except:
				print(f'The template file {tmpltName} does not exist inside the folder {templateFolder}.\nPlease check.\nQuitting...')
				quit()

		for line in tmpltLines:

			if templateEngine == 'textFSM':

				h1 = line.find(DATA_VALUE)
				h2 = line.find(DATA_COMMAND)
				h3 = line.find(DATA_MAJOR_DWN)
				h4 = line.find(DATA_FLTR_COLS)
				h5 = line.find(DATA_FLTR_ACTN)
				h6 = line.find(DATA_SHOW_DIFF_COL)
				
				if h1 != -1:
					# We identify here the variables
					col = line.split(' ')[-2]
					d[tmpltName]['templateColumns'].append(col)
				
				if h2 != -1:
					# Here we identify the command
					line = line.replace(DATA_COMMAND + ' ', DATA_COMMAND)
					cmd  = line.split(':')[1].strip('\n')
					d[tmpltName]['commandKey'] = cmd

				if h3 != -1:
					# we add more major words to the list
					line = line.replace(DATA_MAJOR_DWN + ' ', DATA_MAJOR_DWN)
					keys = line.split(':')[1].strip('\n').split(',')
					for key in keys:
						d[tmpltName]['majorDown'].append(key)

				if h4 != -1:
					# We identify the columnes to be filtered
					line = line.replace(DATA_FLTR_COLS + ' ', DATA_FLTR_COLS)
					keys = line.split(':')[1].strip('\n').split(',')
					for key in keys:
						if key not in [None, '', ' ']:
							d[tmpltName]['filterColumns'].append(key)

				if h5 != -1:
					# we identify the action to be performed on the filterd columns
					line = line.replace(DATA_FLTR_ACTN + ' ', DATA_FLTR_ACTN)
					action = line.split(':')[1].strip('\n')
					d[tmpltName]['filterAction'] = action

				if h6 != -1:
					# we identify which column to add when showing only the differences
					line = line.lstrip().strip('\n')
					keys = line.split(':')[1].strip('\n').split(',')
					for key in keys:
						if key not in [None, '', ' ']:
							key = key.lstrip().rstrip()
							d[tmpltName]['showDiffColumns'].append(key)

			if templateEngine == 'ttp':

				h1 = line.find('#Columns: ')
				h2 = line.find('#Command: ')
				h3 = line.find('#majorDown: ')
				h4 = line.find('#filterColumns: ')
				h5 = line.find('#filterAction: ')
				
				if h1 != -1:
					col = line.split(': ')[1].strip('\n').split(",")
					d[tmpltName]['templateColumns'] = col
				
				if h2 != -1:
					cmd = line.split(': ')[1].strip('\n')
					d[tmpltName]['commandKey'] = cmd

				if h3 != -1:
					keys = line.split(': ')[1].strip('\n').split(',')
					for key in keys:
						d[tmpltName]['majorDown'].append(key)

				if h4 != -1:
					# We identify the columnes to be filtered
					keys = line.split(': ')[1].strip('\n').split(',')
					for key in keys:
						if key not in [None, '', ' ']:
							d[tmpltName]['filterColumns'].append(key)

				if h5 != -1:
					# we identify the action to be performed on the filterd columns
					action = line.split(': ')[1].strip('\n')
					d[tmpltName]['filterAction'] = action

		if len(d[tmpltName]['filterColumns']) > 0:

			print(f'The template {tmpltName} has the following columns to be filtered:')
			print(f'{d[tmpltName]["filterAction"]} the following columns: {d[tmpltName]["filterColumns"]}')

			# checking column's names
			x = [col for col in d[tmpltName]['filterColumns'] if col not in d[tmpltName]['templateColumns']]
			if len(x) > 0:
				print(f'There are some columns which are not original variables of the template.')
				print(x)
				print(f'Check the variables names. Quitting...')
				quit()

			# we want to filter columns
			if d[tmpltName]['filterAction'] not in ['include-only','exclude']:
				# we want to filter columns but we have not specified
				# an action to perform
				print(f'The the action to be used has not been properly set.')
				print(f'Please set either "include-only" or "exclude" in the comments section of the template file.\nQuitting...')
				quit()

			if d[tmpltName]['filterAction'] == 'exclude':

				# we check if the filter columns are equal to the templates' columns.
				# if so, chance is we're getting an empty DF. We don't want this.
				if sorted(d[tmpltName]['filterColumns']) == sorted(d[tmpltName]['templateColumns']):
					print(f'The template {tmpltName} has the following columns:')
					print(d[tmpltName]['templateColumns'])
					print(f'Since the action to be performed is "exclude" all the columns will be filtered out and you')
					print(f'will end with an empty table. Quitting...')
					quit()

				# since we are using 'exclude' our final list of filtered columns is the difference
				# between the original 'templateColumns' and the 'filterColumns'
				x = [col for col in d[tmpltName]['templateColumns'] if col not in d[tmpltName]['filterColumns']]
				d[tmpltName]['filterColumns'] = x

		else:
			# if no filtering columns are defined, we assign those by the original
			# template columns
			d[tmpltName]['filterColumns'] = d[tmpltName]['templateColumns'].copy()

		# We now analyze for showing diff-case columns
		if len(d[tmpltName]['showDiffColumns']) > 0:

			print(f'The template {tmpltName} has the following columns to be shown when displaying diff-results:')
			print(f'columns: {d[tmpltName]["showDiffColumns"]}')

			# checking column's names
			x = [col for col in d[tmpltName]['showDiffColumns'] if col not in d[tmpltName]['templateColumns']]
			if len(x) > 0:
				print(f'There are some columns which are not original variables of the template.')
				print(x)
				print(f'Check the variables names. Quitting...')
				quit()
			
			# If we are filtering columns in the original DataFrame
			# we must be sure that our show-diff columns are not being filtered out...
			if len(d[tmpltName]['filterColumns']) > 0:
				
				diffCols = [x for x in d[tmpltName]['showDiffColumns'] if x not in d[tmpltName]['filterColumns']]
				if len(diffCols) > 0:
					print(f'The template {tmpltName} has the following filtered columns:')
					print(f'{d[tmpltName]["filterAction"]} {d[tmpltName]["filterColumns"]}')
					print(f'The columns you want to use for displaying results are not conisdered inside the filter.')
					print(d[tmpltName]['showDiffColumns'])
					print(f'Quitting...')
					quit()

	print(f'##### Successfully Loaded Templates from folder {templateFolder} #####')
	return d

def makeParsed(nomTemplate, routerLog, templateFolder, templateEngine, templateColumns):
	"""
	Parse through textFSM (reading the file again)

	Args:
		nomTemplate (string): name of file containing the textFSM template
		routerLog (string):   logs of router
		templateFolder (string): folder containing the templates
		templateEngine (string): type of templates
		templateColumns (list): columns in the template

	Returns:
		dataframe with parsed results
	"""

	if templateEngine == 'textFSM':

		if nomTemplate == GENERAL_TEMPL:
			template = io.StringIO(GENERAL_TEMPL_LINES) #Para leer correctamente en textfsm.TextFSM(template)
		else:
			template = open(templateFolder + nomTemplate)

		results_template = textfsm.TextFSM(template)
		parsed_results   = results_template.ParseText(routerLog)

		# With list of results, we build a Pandas DataFrame
		parsed_results = pd.DataFrame(parsed_results, columns= templateColumns)

	if templateEngine == 'ttp':

		with open(templateFolder + nomTemplate) as f:
			template = f.read()

		parser = ttp(data=routerLog, template=template)
		parser.parse()

		output = parser.result(format='table')
		parsed_results = output[0][1][0]

		parsed_results = pd.DataFrame(parsed_results, columns= templateColumns)

	return parsed_results

def readLog(logFolder, formatJson):
	"""
	Reads logs, and stores router logs in memory for processing

	Args:
		logFolder (string):  name of folder
		formatJson (string): "yes" or "no"

	Returns: dictionary with logs
	"""

	if formatJson is True:

		ending = '*rx.json'

	else:

		ending = '*rx.txt'

	if _platform == "linux" or _platform == "linux2" or _platform == "darwin":
		# linux

		listContent  = [f for f in glob.glob(logFolder  + ending)]

	elif _platform == "win64" or _platform == "win32":
		# Windows 64-bit

		listContent  = [f.replace("\\", '/') for f in glob.glob(logFolder  + ending)]
	else:
		print(str(_platform) + ": not a valid platform. Quitting....")
		quit()

	d = {}

	if formatJson is True:

		for name in listContent:
			with open(name) as f:
				d[name] = json.load(f)

	else:
	
		for name in listContent:
			with open(name) as f:
				d[name] = f.read()

	print(f'##### Logs Loaded Successfully from folder {logFolder} #####')

	return d

def parseResults(dTmpl, dLog, templateFolder, templateEngine, routerId):
	"""
	Build the Dataframe from textFSM filter, index and router log

	Args:
		dTmpl (dict):           dictionary with info from templates.
		dLog (dict):            dicitonary with logs. Each key is the fileName; the value, is the content of the log.
		templateFolder (str):   folder of templates
		templateEngine:         textFsm or ttp
		routerId:               name, IP or both

	Returns:
		datosEquipo (dict): Dictionary where keys are templateNames. For each key, a DF with parsed results.
	"""
	
	def detParseStatus(datosCmdsLogs,dfTemp):
		"""
		To determine the parseStatus. Options: no_matching_entries, no_parsing, no_data, ok.
		Here, we don't consider the comparision between pre and post logs (statuses: changes_detected and major_errors)
		"""

		parseStatus = 'no_template'
		if len(datosCmdsLogs) > 0 and len(dfTemp) == 0:
			if any(no_match in datosCmdsLogs for no_match in NO_MATCH):
				parseStatus = 'no_matching_entries'
			else:
				parseStatus = 'no_parsing'
		elif len(dfTemp) == 0 and len(datosCmdsLogs) == 0:
			parseStatus = 'no_data'
		else:
			parseStatus = 'ok'
		
		return parseStatus
	
	def writeDfTemp(dfResult,filterCols,routerId,routerName,routerIP,dfTemp):

		# If there are columns to be filtered, we reduced the 
		# size of the DF to that number of columns
		if len(filterCols) > 0:
			dfResult = dfResult[filterCols]

		# We need to define the identification of the router.
		if 'NAME' in RTR_ID[routerId]:
			dfResult['NAME'] = routerName

		if 'IP' in RTR_ID[routerId]:
			dfResult['IP']   = str(routerIP)

		dfResult = dfResult[orderedColums]
		dfTemp = pd.concat([dfTemp, dfResult])

		# It is stored in the dataEquipment dictionary with the key nomTemplate
		# the DF with the data of all routers
		return dfTemp

	datosEquipo		= {}
	dNoMatchedLog	= {} #Dictionary similar to dLog, but only with noMatched information
	noMatchedCmdAllRtr	= []

	for idR, routerLogKey in enumerate(dLog.keys()): #To each router
		routerLogFname  = routerLogKey.split("/")[-1]

		routerName		= dLog[routerLogKey]['name']
		routerIP		= dLog[routerLogKey]['ip']

		#The previous versions of taskAutom don't have this information
		try:
			routerVersion	= dLog[routerLogKey]['version']
		except:
			routerVersion	= 'NA'
		try:
			routerHwType	= dLog[routerLogKey]['hwType']
		except:
			routerHwType	= 'NA'

		#For use just keys with command:
		command_keys = [k for k in dLog[routerLogKey].keys() if k not in NON_COMMAND_KEYS]
		# logs is each command that was executed in router, inside json file.
		
		noMatchedCmdPerRtr = command_keys.copy()
		# To control which command have no matching template (here, per router)

		#For each command in command_keys(list)
		for cmdsLogs in command_keys: 
			datosCmdsLogs = dLog[routerLogKey][cmdsLogs] #Logs obtained for each command

			#For each template, we test the match with command, 
			for idT, tmpltName in enumerate(dTmpl.keys()):
				commandKey		= dTmpl[tmpltName]['commandKey']

				# command name in each template file
				prog = re.compile(commandKey)

				# searchKey is the regex match between logs and prog
				match = prog.search(cmdsLogs)

				if match and (tmpltName != GENERAL_TEMPL):
					# If there's a match, we take that command off the list noMatchedCmdPerRtr
					# Important for processing cases that use generic template
					noMatchedCmdPerRtr.remove(cmdsLogs)
					
					# if command(in template) == command(in key of router) then we stores log info in routeLog variable
					# Means that the command was executed and there's a template to this command

					# {
					# 	'logs1':'output1',
					# 	'logs2':'output2',
					# 	'logsN':'outputN',
					# }

					# "/show router 4001 route-table | match No": "No. of Routes: 566",
					# "/show router 4002 route-table | match No": "MINOR: CLI Invalid router \"4002\".\u0007",
					# "/show router route-table | match No": "No. of Routes: 3337",

					templateColumns	= dTmpl[tmpltName]['templateColumns']
					filterCols		= dTmpl[tmpltName]['filterColumns']
					orderedColums	= RTR_ID[routerId] + filterCols

					if tmpltName not in datosEquipo:
						datosEquipo[tmpltName] = {}

					datosEquipo[tmpltName]['command']		= cmdsLogs
					datosEquipo[tmpltName]['template']		= tmpltName

					if 'dfResultDatos' not in datosEquipo[tmpltName]:
						datosEquipo[tmpltName]['dfResultDatos'] = pd.DataFrame()

					routerLog = cmdsLogs + '\n' + datosCmdsLogs + '\n' #Command and your data

					# We parse results from the key:value association
					# A list is returnd with results
					# to parse, with provide the complete set of columns as defined inside the template: templateColumns
					dfResult								= makeParsed(tmpltName, routerLog, templateFolder, templateEngine, templateColumns)
					datosEquipo[tmpltName]['dfResultDatos']	= writeDfTemp(dfResult, filterCols,routerId,routerName,routerIP, datosEquipo[tmpltName]['dfResultDatos'])
					datosEquipo[tmpltName]['parseStatus']	= detParseStatus(datosCmdsLogs, datosEquipo[tmpltName]['dfResultDatos'])

		# Writing dNoMatchedLog for each router. At this point, all attempts to match templates already occurred.
		for cmdsLogs in noMatchedCmdPerRtr:

			if cmdsLogs not in noMatchedCmdAllRtr:
				noMatchedCmdAllRtr.append(cmdsLogs) 
				# Adding to list to no-matched commands, containing information of all routers
			
			#Adding the information to dNoMatchedLog: the dict that only has the cases of no-matched command
			if routerLogKey not in dNoMatchedLog:
				dNoMatchedLog[routerLogKey] = {}

			#This information is required in the following "for"
			dNoMatchedLog[routerLogKey][cmdsLogs]= dLog[routerLogKey][cmdsLogs] 
			dNoMatchedLog[routerLogKey]['ip']	= dLog[routerLogKey]['ip']		# for writeDfTemp
			dNoMatchedLog[routerLogKey]['name']	= dLog[routerLogKey]['name']	# for writeDfTemp

	#Processing the no-matched commands
	for idR, routerLogKey in enumerate(dNoMatchedLog.keys()):
		# Basic information of the router, for writeDfTemp
		routerName	= dLog[routerLogKey]['name']
		routerIP	= dLog[routerLogKey]['ip']

		for i,cmdsLogs in enumerate(noMatchedCmdAllRtr):
			# Here, enumerate and noMatchedCmdAllRtr are used to ensure that all no-matched commands,
			# indenpendently of the router, will have the same identification (i).
			# So we can keep the information of the same command together correctly

			# If certain router have datosCmdsLogs, we use this information in this iteration.
			# Otherwise, moves to the other iteration of "for" of commands.
			try:
				datosCmdsLogs = dNoMatchedLog[routerLogKey][cmdsLogs] #Logs obtained for each command
			except:
				continue

			templateColumns	= dTmpl[GENERAL_TEMPL]['templateColumns']
			filterCols		= dTmpl[GENERAL_TEMPL]['filterColumns']
			orderedColums	= RTR_ID[routerId] + filterCols

			if f'general_{i}' not in datosEquipo:
				datosEquipo[f'general_{i}'] = {}

			datosEquipo[f'general_{i}']['command']	= cmdsLogs
			datosEquipo[f'general_{i}']['template']	= GENERAL_TEMPL

			if 'dfResultDatos' not in datosEquipo[f'general_{i}']:
				datosEquipo[f'general_{i}']['dfResultDatos'] = pd.DataFrame()

			routerLog = cmdsLogs + '\n' + datosCmdsLogs + '\n'

			dfResult									 = makeParsed(GENERAL_TEMPL, routerLog, '', templateEngine, templateColumns)
			datosEquipo[f'general_{i}']['dfResultDatos'] = writeDfTemp(dfResult, filterCols,routerId,routerName,routerIP, datosEquipo[f'general_{i}']['dfResultDatos'])
			datosEquipo[f'general_{i}']['parseStatus']	 = detParseStatus(datosCmdsLogs, datosEquipo[f'general_{i}']['dfResultDatos'])

	return datosEquipo

def searchDiffAll(datosEquipoPre, datosEquipoPost, dTmplt, routerId):
	'''
	Makes a new table, in which it brings the differences between two tables (post-pre)
	'''

	countDif   = {}

	for tmpltName in datosEquipoPre.keys():
		if tmpltName not in countDif:
			countDif[tmpltName] = {}

		# En datosEquipoPre[tmpltName], tmpltName puede ser generic_0 y esa key no existe en lo dict de templates (dTmplt).
		# Por eso que es necessario gravar el template en datosEquipoPre[tmpltName]['template'], para entonces relacionar ese
		# valor con el dTmplt[template] (en los ejemplos de template generico, template == 'generic.template')

		template = datosEquipoPre[tmpltName]['template']
		filterCols = dTmplt[template]['filterColumns']
		dfUnion = pd.merge(datosEquipoPre[tmpltName]['dfResultDatos'], datosEquipoPost[tmpltName]['dfResultDatos'], how='outer', indicator='Where').drop_duplicates()
		dfInter = dfUnion[dfUnion.Where=='both']
		dfCompl = dfUnion[~(dfUnion.isin(dfInter))].dropna(axis=0, how='all').drop_duplicates()
		dfCompl['Where'] = dfCompl['Where'].str.replace('left_only','Pre')
		dfCompl['Where'] = dfCompl['Where'].str.replace('right_only','Post')

		orderedColums = RTR_ID[routerId] + filterCols

		countDif[tmpltName]['dfResultDatos'] = dfCompl.sort_values(by = orderedColums)

	return countDif

def searchDiffOnly(datosEquipoPre, datosEquipoPost, dTmplt, routerId):
	'''
	Makes a new table, in which it brings just the differences between two tables (post-pre)
	'''

	countDif   = {}

	for tmpltName in datosEquipoPre.keys():
		if tmpltName not in countDif:
			countDif[tmpltName] = {}

		template = datosEquipoPre[tmpltName]['template']
		filterCols = dTmplt[template]['showDiffColumns']

		dfPre  = datosEquipoPre[tmpltName]['dfResultDatos']
		dfPost = datosEquipoPost[tmpltName]['dfResultDatos']

		dfMerge = pd.merge(dfPre,dfPost, suffixes=('_pre','_post'), how='outer', indicator='Where')
		dfMerge['Where'] = dfMerge['Where'].str.replace('left_only','Pre')
		dfMerge['Where'] = dfMerge['Where'].str.replace('right_only','Post')
		dfMerge          = dfMerge[dfMerge['Where'].isin(['Pre','Post'])]

		routers = dfMerge['NAME'].unique()

		dfMerge_new  = pd.DataFrame()

		for router in routers:

			dfRouter = pd.DataFrame()

			tempMerge = dfMerge[dfMerge['NAME']==router]

			if len(tempMerge) > 1:
				tempMerge = tempMerge.loc[:,tempMerge.nunique() > 1]
				tempMerge['NAME'] = router

			dfRouter = pd.concat([dfRouter,tempMerge])

			routerCols   = list(dfRouter.columns)
			dfRouter     = pd.merge(dfRouter,dfMerge, on=routerCols)
			hasFilterCol = len([x for x in routerCols if x in filterCols])

			if hasFilterCol == 0:
				dfRouter   = dfRouter[routerCols + filterCols]
			else:
				dfRouter   = dfRouter[routerCols]

			dfMerge_new = pd.concat([dfMerge_new,dfRouter])

		dfMerge_new.reset_index(inplace=True)
		dfMerge_new = dfMerge_new.drop(columns='index')

		if len(dfMerge_new) > 0:

			finalColumns = list(dfMerge_new.columns)
			finalColumns.remove('NAME')
			finalColumns.remove('Where')

			if len(filterCols) > 0:
				[finalColumns.remove(x) for x in filterCols]
				finalColumns = ['NAME'] + filterCols + finalColumns + ['Where']
			else:
				finalColumns = ['NAME'] + finalColumns + ['Where']

			dfMerge_new         = dfMerge_new[finalColumns]
			countDif[tmpltName]['dfResultDatos'] = dfMerge_new.sort_values(by = finalColumns)
		else:
			countDif[tmpltName]['dfResultDatos'] = dfMerge_new

	return countDif


def findMajor(count_dif, dTmplt, routerId, showResults, datosEquipoPre):
	'''
	Makes a table from the results of searching for Major errors in the post table define in yml file for specific template,\n
	or down if is not define the words for the template, which are not in the Pre table
	'''

	countDown  = {}
	
	for tmpltName in count_dif.keys():
		if tmpltName not in countDown:
			countDown[tmpltName] = {}

		df         = pd.DataFrame()
		template = datosEquipoPre[tmpltName]['template']

		for majorWord in dTmplt[template]['majorDown']:

			filterCols = dTmplt[template]['filterColumns']

			if 'Where' in count_dif[tmpltName]['dfResultDatos'].columns:

				df1 = count_dif[tmpltName]['dfResultDatos'][count_dif[tmpltName]['dfResultDatos']['Where']=='Post']
				
				if len(df1) > 0:
					df1 = df1[df1.apply(lambda r: r.str.contains(majorWord, case=False).any(), axis=1)]
				else:
					df1 = pd.DataFrame(columns=count_dif[tmpltName]['dfResultDatos'].columns)

				df = pd.concat([df, df1])

				if showResults == 'all':
					df = df.sort_values(by = RTR_ID[routerId] + filterCols)

		countDown[tmpltName]['dfResultDatos'] = df

	return countDown

def makeTable(datosEquipoPre, datosEquipoPost):
	'''
	Sort the table pre and post to present in Excel
	'''

	df_all          = {}
	datosEquipoPre1 = datosEquipoPre.copy()
	
	for tmpltName in datosEquipoPre.keys():
		if tmpltName not in df_all:
				df_all[tmpltName] = {}

		datosEquipoPre1[tmpltName]['##']='##'

		df_all[tmpltName]['dfResultDatos']	= pd.concat([datosEquipoPre1[tmpltName]['dfResultDatos'], datosEquipoPost[tmpltName]['dfResultDatos']], axis=1, keys=('Pre-Check', 'Post-Check'))
		df_all[tmpltName]['parseStatus']	= datosEquipoPost[tmpltName]['parseStatus']
		df_all[tmpltName]['command']		= datosEquipoPre[tmpltName]['command']

	return df_all

def constructExcel(df_final, count_dif, searchMajor, folderLog):
	"""
	Sort the data and format creating the Excel
	_summary_

	Args:
		df_final (pandas): DataFrame with pre and post data
		count_dif (pandas): DataFrame with only differences
		searchMajor (pandas): DataFrame with only errors
		folderLog (string): name of the folder
	"""

	fileName  = folderLog[:-1] + ".xlsx"

	writer    = pd.ExcelWriter(fileName, engine='xlsxwriter') #creates instance of an excel workbook
	workbook  = writer.book

	# Create index tab
	indexSheet = workbook.add_worksheet('index')

	print('\nSaving Excel')

	for idx,template in enumerate(df_final.keys()):

		dfData  = df_final[template]['dfResultDatos']
		dfDiff  = count_dif[template]['dfResultDatos']
		dfMajor = searchMajor[template]['dfResultDatos']
		dfParseStatus = df_final[template]['parseStatus']

		sheet_name = template.replace('nokia_sros_','')
		sheet_name = sheet_name.replace('.template','')
		sheet_name = sheet_name.replace('_template','')
		sheet_name = sheet_name.replace('.ttp','')
		sheet_name = sheet_name.replace('.','_')

		if len(sheet_name) > 31:
			sheet_name = sheet_name[:31]

		# Selecting Tab's color and error messages
		if dfParseStatus not in ['ok','changes_detected','major_errors']:
			output = dfParseStatus
		elif len(dfMajor) == 0 and len(dfDiff) == 0:
			output = 'ok'
		elif len(dfMajor) == 0 and len(dfDiff) != 0:
			output = 'changes_detected'
		elif len(dfMajor) != 0:
			output = 'major_errors'

		# cell format
		cell_format  = workbook.add_format({'color': CELL_COLOR, 'font_size': CELL_FONT_SIZE, 'fg_color': D_STATUS[output]['colorTab'], 'align': 'center', 'border': 1 ,'bold': True})

		srcCol   = 'A'+str(idx+1)

		# Building index
		for k, i_dict in INDEX_COL.items():
			indexSheet.write(0,i_dict['position'],i_dict['colName'], workbook.add_format({'font_size':CELL_FONT_SIZE,'align':'center','border':1,'bold':True}))
			indexSheet.set_column(i_dict['col'],i_dict['width'])

		indexSheet.write_url(idx+1,0, 'internal:'+sheet_name+'!A1', string=sheet_name)
		indexSheet.write(idx+1,1, df_final[template]['command'])
		indexSheet.write(idx+1,2, D_STATUS[output]['shortText'], cell_format)

		# Creating Tab
		worksheet = workbook.add_worksheet(sheet_name)
		worksheet.set_tab_color(D_STATUS[output]['colorTab'])
		writer.sheets[sheet_name] = worksheet
		dfData.to_excel(writer, sheet_name=sheet_name, startrow=0, startcol=0) #Creates Excel File
		worksheet.write_url('A1', 'internal:index!A1', string='Index')
		
		if len(dfDiff) > 0:
			### Changes Section
			srcCol   = 'A'+str(len(dfData)+5)
			dstCol   = 'H'+str(len(dfData)+5)
			colRange = srcCol + ':' + dstCol
			warnTex  = D_STATUS[output]['warnText']
			worksheet.merge_range(colRange, warnTex, cell_format)
			if len(dfDiff) > 0:
				dfDiff.to_excel(writer, sheet_name=sheet_name, startrow=len(dfData)+6, startcol=0)
		if len(dfMajor) > 0:
			### Major Error Section
			srcCol   = 'A'+str((len(dfData)+(len(dfDiff)))+9)
			dstCol   = 'H'+str((len(dfData)+(len(dfDiff)))+9)
			colRange = srcCol + ':' + dstCol
			errText   = warnTex  = D_STATUS[output]['errText']
			worksheet.merge_range(colRange, errText, cell_format)
			if len(dfMajor) > 0:
				dfMajor.to_excel(writer, sheet_name=sheet_name, startrow=(len(dfData)+(len(dfDiff)))+10, startcol=0)

		print('#',idx,template)
	
	writer.close() #saves workbook to file in python file directory

def renderAtp(dictParam):
	"""[Generates a ATP based on the json logs obtained from taskAutom.]

	Args:
		dictParam

	Returns:
		None
	"""

	preFolder  = dictParam['preFolder']
	postFolder = dictParam['postFolder']

	jsonFilesPre = [preFolder+x for x in os.listdir(preFolder) if '.json' in x and x != '00_report.json']
	if postFolder != '':
		jsonFilesPos = [postFolder+x for x in os.listdir(postFolder) if '.json' in x and x != '00_report.json']

	job0docx = "./ATP.docx"

	print("\nGenerating ATP: " + job0docx)

	myDoc    = docx.Document()
	myStyles = myDoc.styles

	styleConsole = myStyles.add_style('Console', WD_STYLE_TYPE.PARAGRAPH)
	styleConsole.font.name = 'Courier'
	styleConsole.font.size = Pt(9)
	styleConsole.paragraph_format.keep_together = True

	styleConsole.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
	#styleConsole.paragraph_format.line_spacing = Pt(10)
	#styleConsole.paragraph_format.line_spacing = .2
	styleConsole.paragraph_format.space_after = Pt(2)

	myDoc.add_heading('ATP', 0)

	charset_allowed = [chr(c) for c in range(32,127)] + ['\n']

	if preFolder != '':

		docMainTitle = myDoc.add_paragraph('Pre-Check')
		docMainTitle.style = myDoc.styles['Heading 1']
		docMainTitle.paragraph_format.line_spacing = 1.5

		for f in jsonFilesPre:
			
			with open(f) as myFile:
				
				logs  = json.load(myFile)
				keys  = [x for x in logs.keys() if 'show' in x]

				routerTitle = f'Router {logs["name"]} ({logs["ip"]})'

				docRouterTitle = myDoc.add_paragraph(routerTitle)
				docRouterTitle.style = myDoc.styles['Heading 2']
				docRouterTitle.paragraph_format.line_spacing = 1.5

				for key in keys:
					showTitle   = key.rstrip('\n').lstrip('\n')
					showContent = ''.join([x for x in logs[key] if x in charset_allowed]).rstrip('\n').lstrip('\n')

					docShowTitle = myDoc.add_paragraph(showTitle)
					docShowTitle.style = myDoc.styles['Heading 3']
					docShowTitle.paragraph_format.line_spacing = 1.5

					docShowContent = myDoc.add_paragraph(showContent)
					docShowContent.style = myDoc.styles['Console']

	if postFolder != '':

		docMainTitle = myDoc.add_paragraph('Post-Check')
		docMainTitle.style = myDoc.styles['Heading 1']
		docMainTitle.paragraph_format.line_spacing = 1.5

		for f in jsonFilesPos:
			
			with open(f) as myFile:
				
				logs  = json.load(myFile)
				keys  = [x for x in logs.keys() if 'show' in x]

				routerTitle = f'Router {logs["name"]} ({logs["ip"]})'

				docRouterTitle = myDoc.add_paragraph(routerTitle)
				docRouterTitle.style = myDoc.styles['Heading 2']
				docRouterTitle.paragraph_format.line_spacing = 1.5

				for key in keys:
					showTitle   = key.rstrip('\n').lstrip('\n')
					showContent = ''.join([x for x in logs[key] if x in charset_allowed]).rstrip('\n').lstrip('\n')

					docShowTitle = myDoc.add_paragraph(showTitle)
					docShowTitle.style = myDoc.styles['Heading 3']
					docShowTitle.paragraph_format.line_spacing = 1.5

					docShowContent = myDoc.add_paragraph(showContent)
					docShowContent.style = myDoc.styles['Console']

	myDoc.save(job0docx)

	print("ATP done...")

def fncRun(dictParam):

	preFolder          = dictParam['preFolder']
	postFolder         = dictParam['postFolder']
	csvTemplate        = dictParam['csvTemplate']
	formatJson         = dictParam['formatJson']
	templateFolder     = dictParam['templateFolder']
	templateEngine     = dictParam['templateEngine']
	templateFolderPost = dictParam['templateFolderPost']
	routerId           = dictParam['routerId']
	showResults        = dictParam['showResults']
	genAtp             = dictParam['genAtp']

	if _platform == "win64" or _platform == "win32":
		templateFolder = templateFolder.replace('/', '\\')
		if templateFolderPost != '':
			templateFolderPost = templateFolderPost.replace('/','\\')

	if preFolder != '' and postFolder == '':

		dTmplt = readTemplate(csvTemplate, templateFolder, templateEngine)
		dLog   = readLog(preFolder, formatJson)

		df_final    = parseResults(dTmplt, dLog, templateFolder, templateEngine, routerId)
		count_dif   = {}
		searchMajor = {}

		for tmpltName in df_final.keys():
			if tmpltName not in count_dif:
				count_dif[tmpltName] = {}
			count_dif[tmpltName]['dfResultDatos']   = pd.DataFrame(columns=df_final[tmpltName]['dfResultDatos'].columns)
			if tmpltName not in searchMajor:
				searchMajor[tmpltName] = {}
			searchMajor[tmpltName]['dfResultDatos'] = pd.DataFrame(columns=df_final[tmpltName]['dfResultDatos'].columns)

		constructExcel(df_final, count_dif, searchMajor, preFolder)

		if genAtp is True:
			renderAtp(dictParam)

	elif preFolder != '' and postFolder != '':

		if templateFolder == templateFolderPost:
			dTmpltPre  = readTemplate(csvTemplate, templateFolder, templateEngine)
			dTmpltPost = readTemplate(csvTemplate, templateFolderPost, templateEngine)
		elif templateFolder != '' and templateFolderPost == '':
			templateFolderPost = templateFolder
			dTmpltPre  = readTemplate(csvTemplate, templateFolder, templateEngine)
			dTmpltPost = readTemplate(csvTemplate, templateFolderPost, templateEngine)
		else:
			dTmpltPre  = readTemplate(csvTemplate, templateFolder, templateEngine)
			dTmpltPost = readTemplate(csvTemplate, templateFolderPost, templateEngine)
			keysPre    = sorted(list(dTmpltPre.keys()))
			keysPos    = sorted(list(dTmpltPost.keys()))

			if keysPre == keysPos:
				pass
			# else:
			# 	if csvTemplate == '':
			# 		if len(keysPre) != len(keysPos):
			# 			print(f'The PRE template folder, {templateFolder}, has {len(keysPre)} templates.')
			# 			print(f'The POST template folder, {templateFolderPost}, has {len(keysPos)} templates.')
			# 			print('Make sure the amount of templates in each folder, is the same. Or use a CSV list of templates.\nQuitting...')
			# 			quit()
			# 		else:
			# 			print(f'The template folders {templateFolder} and {templateFolderPost} have the same amount of templates')
			# 			print('But there are differences among them.')
			# 			print('Check the contents. Quitting...')
			# 			quit()
			# 	else:
			# 		pass

		dLogPre  = readLog(preFolder, formatJson)
		dLogPost = readLog(postFolder, formatJson)
			
		datosEquipoPre  = parseResults(dTmpltPre,  dLogPre,  templateFolder,     templateEngine, routerId)
		datosEquipoPost = parseResults(dTmpltPost, dLogPost, templateFolderPost, templateEngine, routerId)

		if showResults == 'all':
			count_dif       = searchDiffAll(datosEquipoPre, datosEquipoPost, dTmpltPre, routerId)
		else:
			count_dif       = searchDiffOnly(datosEquipoPre, datosEquipoPost, dTmpltPre, routerId)

		searchMajor     = findMajor(count_dif, dTmpltPre, routerId, showResults, datosEquipoPre)
		df_final        = makeTable(datosEquipoPre, datosEquipoPost)

		constructExcel(df_final, count_dif, searchMajor, postFolder)

		if genAtp is True:
			renderAtp(dictParam)

	elif preFolder == '':
		print('No PRE folder defined. Please Verify.')



def main():

	parser1 = argparse.ArgumentParser(description='Log Analysis', prog='PROG', usage='%(prog)s [options]')
	parser1.add_argument('-pre', '--preFolder',     type=str, required=True, help='Folder with PRE Logs. Must end in "/"',)
	parser1.add_argument('-post','--postFolder' ,   type=str, default='',    help='Folder with POST Logs. Must end in "/"',)
	parser1.add_argument('-csv', '--csvTemplate',   type=str, default='', help='CSV with list of templates names to be used in parsing. If the file is omitted, then all the templates inside --templateFolder, will be considered for parsing. Default=None.')
	parser1.add_argument('-json', '--formatJson',   type=str, default = 'yes', choices=['yes','no'], help='logs in json format: yes or no. Default=yes.')
	parser1.add_argument('-tf', '--templateFolder', type=str, default='Templates/', help='Folder where templates reside. Used both for PRE and POST logs. Default=Templates/')
	parser1.add_argument('-tf-post', '--templateFolderPost', type=str, default='', help='If set, use this folder of templates for POST logs.')
	parser1.add_argument('-te', '--templateEngine', choices=['ttp','textFSM'], default='textFSM', type=str, help='Engine for parsing. Default=textFSM.')
	parser1.add_argument('-ri', '--routerId',       choices=['name','ip','both'], default='name', type=str, help='Router Id to be used within the tables in the Excel report. Default=name.')
	parser1.add_argument('-sr', '--showResults',    choices=['all','diff'], default='all', type=str, help='When comparison is done, show all variables or only the differences. Only available if --ri/--routerId=name. Default=all)')
	parser1.add_argument('-ga', '--genAtp',        type=str, help='Generate ATP document in docx format, based on the contents of the json files from taskAutom. Default=no', default='no', choices=['no','yes'])
	parser1.add_argument('-v'  ,'--version',        help='Version', action='version', version='(c) 2024 - Version: 4.2.2' )

	args               = parser1.parse_args()

	dictParam = dict(
		preFolder          = args.preFolder,
		postFolder         = args.postFolder,
		csvTemplate        = args.csvTemplate,
		formatJson         = True if args.formatJson == 'yes' else False,
		templateFolder     = args.templateFolder,
		templateEngine     = args.templateEngine,
		templateFolderPost = args.templateFolderPost,
		routerId           = args.routerId,
		showResults        = args.showResults,
		genAtp             = True if args.genAtp == 'yes' else False,
	)

	if dictParam['showResults'] == 'diff' and dictParam['routerId'] != 'name':
		print('If showResults is "diff", routerId must be "name"\nQuitting ...')
		quit()

	fncRun(dictParam)

### To be run from the python shell
if __name__ == '__main__':
	main()